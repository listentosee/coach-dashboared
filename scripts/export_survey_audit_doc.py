#!/usr/bin/env python3

from __future__ import annotations

import argparse
import html
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document


SKIP_WIDGET_TYPES = {"Text", "Paragraph", "Button", "ThankYou"}
OPEN_ENTRY_WIDGET_TYPES = {"LongAnswer", "NumberInput"}
CHOICE_WIDGET_TYPES = {"Checkboxes", "MultipleChoice", "ImagePicker"}


@dataclass
class Question:
    heading: str
    title: str
    caption: str
    question_type: str
    distractors: str


@dataclass
class Group:
    name: str
    questions: list[Question]


def clean_html_text(value: str | None) -> str:
    if not value:
        return ""

    text = value
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</(p|div|h[1-6]|ul|ol)>", "\n", text)
    text = re.sub(r"(?i)<li[^>]*>", "- ", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text).replace("\xa0", " ")

    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def get_logic_value(container: dict, key: str) -> str:
    field = container.get(key, {})
    if not isinstance(field, dict):
        return ""

    logic = field.get("logic", {})
    if isinstance(logic, dict):
        return clean_html_text(logic.get("value"))
    if isinstance(logic, str):
        return clean_html_text(logic)
    return ""


def get_widget_text(widget: dict, key: str) -> str:
    return get_logic_value(widget.get("template", {}), key)


def extract_static_options(widget: dict) -> list[str]:
    options = widget.get("template", {}).get("options", {}).get("staticOptions", [])
    values: list[str] = []
    for option in options:
        value = get_logic_value(option, "value") or get_logic_value(option, "label")
        if value:
            values.append(value)
    return values


def extract_distractors(widget: dict, widget_type: str) -> str:
    static_options = extract_static_options(widget)
    if widget_type in CHOICE_WIDGET_TYPES and static_options:
        return "; ".join(static_options)

    if widget_type == "OpinionScale":
        template = widget.get("template", {})
        min_value = template.get("minValue")
        max_value = template.get("maxValue")
        left_label = get_widget_text(widget, "leftLabel")
        right_label = get_widget_text(widget, "rightLabel")

        scale_parts: list[str] = []
        if min_value is not None and max_value is not None:
            scale_parts.append(f"{min_value}-{max_value}")
        if left_label or right_label:
            scale_parts.append(f"({left_label or '?'} -> {right_label or '?'})")

        return " ".join(scale_parts) if scale_parts else "None"

    if widget_type in OPEN_ENTRY_WIDGET_TYPES:
        return "None"

    if static_options:
        return "; ".join(static_options)

    return "None"


def resolve_question_title(widget: dict, fallback_index: int) -> str:
    title = get_widget_text(widget, "label")
    if title:
        return title

    fallback_name = widget.get("name", "").strip()
    if fallback_name and not fallback_name.startswith("Untitled "):
        return fallback_name

    return f"Question {fallback_index}"


def extract_groups(path: Path) -> list[Group]:
    data = json.loads(path.read_text())
    steps = data.get("template", {}).get("steps", {})
    groups: list[Group] = []

    for step in steps.values():
        group_name = step.get("name", "").strip() or "Untitled Section"
        widgets = step.get("template", {}).get("widgets", {})
        ordered_widgets = sorted(
            widgets.values(),
            key=lambda widget: (
                widget.get("position", {}).get("row", 0),
                widget.get("position", {}).get("column", 0),
                widget.get("id", ""),
            ),
        )

        questions: list[Question] = []
        for widget in ordered_widgets:
            widget_type = widget.get("type", "").strip() or "Unknown"
            if widget_type in SKIP_WIDGET_TYPES:
                continue

            title = resolve_question_title(widget, len(questions) + 1)
            questions.append(
                Question(
                    heading=title,
                    title=title,
                    caption=get_widget_text(widget, "caption") or "None",
                    question_type=widget_type,
                    distractors=extract_distractors(widget, widget_type),
                )
            )

        if questions:
            groups.append(Group(name=group_name, questions=questions))

    return groups


def add_labeled_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value)


def build_document(survey_name: str, groups: list[Group], output_path: Path) -> None:
    document = Document()
    document.add_heading(survey_name, level=0)

    if not groups:
        document.add_paragraph("No interactive questions found.")
    else:
        for group in groups:
            document.add_heading(group.name, level=2)
            for question in group.questions:
                document.add_heading(question.heading, level=3)
                add_labeled_paragraph(document, "Title", question.title)
                add_labeled_paragraph(document, "Caption", question.caption)
                add_labeled_paragraph(document, "Type", question.question_type)
                add_labeled_paragraph(document, "Distractors (values)", question.distractors)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_output_path(output_dir: Path, survey_path: Path) -> Path:
    return output_dir / f"{survey_path.stem} audit extraction.docx"


def main() -> None:
    parser = argparse.ArgumentParser(description="Export survey questions to Word documents.")
    parser.add_argument(
        "--input-dir",
        type=Path,
        default=Path("docs/Surveys"),
        help="Directory containing survey JSON files.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        help="Directory to write one .docx file per survey. Defaults to the input directory.",
    )
    args = parser.parse_args()

    input_dir = args.input_dir
    output_dir = args.output_dir or input_dir
    survey_files = sorted(input_dir.glob("*.json"))

    for survey_path in survey_files:
        groups = extract_groups(survey_path)
        build_document(survey_path.stem, groups, build_output_path(output_dir, survey_path))


if __name__ == "__main__":
    main()
